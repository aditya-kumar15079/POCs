"""
Word Document Processor for TestFoundry Framework
Handles Word document parsing with text, image, and table extraction
"""

import io
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pytesseract
from PIL import Image

class WordProcessor:
    """Word document processor with comprehensive content extraction"""
    
    def __init__(self, config: Dict[str, Any], logger: logging.Logger):
        """Initialize Word processor
        
        Args:
            config: Configuration dictionary
            logger: Logger instance
        """
        self.config = config
        self.logger = logger
        self.word_config = config.get('document_processing', {}).get('word', {})
        
        # Settings
        self.extract_images = self.word_config.get('extract_images', True)
        self.extract_tables = self.word_config.get('extract_tables', True)
        self.preserve_formatting = self.word_config.get('preserve_formatting', True)
    
    async def process_document(self, file_path: Path) -> Dict[str, Any]:
        """Process Word document and extract content
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Dictionary containing extracted content
        """
        try:
            self.logger.info(f"Processing Word document: {file_path.name}")
            
            # Load document
            doc = Document(str(file_path))
            
            # Extract content
            pages_content = await self._extract_content(doc)
            
            # Extract images if enabled
            images = []
            if self.extract_images:
                images = await self._extract_images(doc)
            
            # Extract tables if enabled
            tables = []
            if self.extract_tables:
                tables = await self._extract_tables(doc)
            
            # Organize content by pages (Word doesn't have explicit pages, so we simulate)
            pages = await self._organize_into_pages(pages_content, images, tables)
            
            # Create document structure
            doc_content = {
                'name': file_path.stem,
                'file_path': str(file_path),
                'file_type': 'word',
                'processor': 'WordProcessor',
                'pages': pages
            }
            
            # Calculate statistics
            total_text = sum(len(page.get('content', '')) for page in pages)
            doc_content['statistics'] = {
                'total_pages': len(pages),
                'total_paragraphs': len(pages_content.get('paragraphs', [])),
                'total_characters': total_text,
                'total_words': len(' '.join(page.get('content', '') for page in pages).split()),
                'images_extracted': len(images),
                'tables_extracted': len(tables)
            }
            
            self.logger.info(f"Word processing complete: {doc_content['statistics']['total_pages']} pages, "
                           f"{doc_content['statistics']['total_characters']} characters")
            
            return doc_content
            
        except Exception as e:
            self.logger.error(f"Error processing Word document {file_path.name}: {e}")
            return {
                'name': file_path.stem,
                'file_path': str(file_path),
                'file_type': 'word',
                'pages': [],
                'error': str(e)
            }
    
    async def _extract_content(self, doc: DocumentType) -> Dict[str, Any]:
        """Extract text content from Word document
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Dictionary containing extracted content
        """
        content = {
            'paragraphs': [],
            'headers': [],
            'footers': []
        }
        
        try:
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_info = {
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else None,
                        'alignment': str(paragraph.alignment) if paragraph.alignment else None,
                        'is_heading': self._is_heading_style(paragraph.style.name if paragraph.style else "")
                    }
                    content['paragraphs'].append(para_info)
            
            # Extract headers
            for section in doc.sections:
                if section.header:
                    header_text = ""
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_text += paragraph.text + "\n"
                    if header_text.strip():
                        content['headers'].append(header_text.strip())
                
                # Extract footers
                if section.footer:
                    footer_text = ""
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text += paragraph.text + "\n"
                    if footer_text.strip():
                        content['footers'].append(footer_text.strip())
            
        except Exception as e:
            self.logger.warning(f"Error extracting text content: {e}")
        
        return content
    
    async def _extract_images(self, doc: DocumentType) -> List[Dict[str, Any]]:
        """Extract images from Word document
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of image information dictionaries
        """
        images = []
        
        try:
            # Get all relationships that are images
            image_rels = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_rels.append(rel)
            
            for idx, rel in enumerate(image_rels):
                try:
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # OCR on image if enabled
                    text_content = ""
                    try:
                        image_pil = Image.open(io.BytesIO(image_data))
                        text_content = pytesseract.image_to_string(image_pil)
                    except Exception:
                        pass
                    
                    image_info = {
                        'index': idx,
                        'filename': rel.target_ref.split('/')[-1],
                        'text_content': text_content.strip(),
                        'size_bytes': len(image_data),
                        'content_type': rel.target_part.content_type
                    }
                    
                    images.append(image_info)
                    
                except Exception as e:
                    self.logger.warning(f"Error extracting image {idx}: {e}")
                    continue
                    
        except Exception as e:
            self.logger.warning(f"Error extracting images: {e}")
        
        return images
    
    async def _extract_tables(self, doc: DocumentType) -> List[Dict[str, Any]]:
        """Extract tables from Word document
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of table information dictionaries
        """
        tables = []
        
        try:
            for table_idx, table in enumerate(doc.tables):
                try:
                    # Extract table data
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        table_data.append(row_data)
                    
                    # Convert table to text representation
                    table_text = ""
                    for row in table_data:
                        table_text += " | ".join(row) + "\n"
                    
                    table_info = {
                        'index': table_idx,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0,
                        'content': table_text.strip(),
                        'data': table_data
                    }
                    
                    tables.append(table_info)
                    
                except Exception as e:
                    self.logger.warning(f"Error extracting table {table_idx}: {e}")
                    continue
                    
        except Exception as e:
            self.logger.warning(f"Error extracting tables: {e}")
        
        return tables
    
    async def _organize_into_pages(self, 
                                 content: Dict[str, Any], 
                                 images: List[Dict[str, Any]], 
                                 tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Organize content into simulated pages
        
        Args:
            content: Extracted text content
            images: Extracted images
            tables: Extracted tables
            
        Returns:
            List of page dictionaries
        """
        pages = []
        
        # Estimate characters per page (typical Word document)
        chars_per_page = 2500
        
        # Combine all text content
        all_text = ""
        paragraphs = content.get('paragraphs', [])
        
        # Add headers and footers
        for header in content.get('headers', []):
            all_text += header + "\n\n"
        
        # Add paragraph content
        for para in paragraphs:
            all_text += para['text'] + "\n\n"
        
        # Add footer content
        for footer in content.get('footers', []):
            all_text += footer + "\n\n"
        
        # Split into pages
        if len(all_text) <= chars_per_page:
            # Single page
            pages.append({
                'page_number': 1,
                'content': all_text.strip(),
                'images': images,
                'tables': tables,
                'has_text': len(all_text.strip()) > 0,
                'character_count': len(all_text)
            })
        else:
            # Multiple pages
            current_page = 1
            current_content = ""
            current_images = []
            current_tables = []
            
            # Distribute paragraphs across pages
            for para in paragraphs:
                para_text = para['text'] + "\n\n"
                
                if len(current_content) + len(para_text) > chars_per_page and current_content:
                    # Create page
                    pages.append({
                        'page_number': current_page,
                        'content': current_content.strip(),
                        'images': current_images,
                        'tables': current_tables,
                        'has_text': len(current_content.strip()) > 0,
                        'character_count': len(current_content)
                    })
                    
                    # Start new page
                    current_page += 1
                    current_content = para_text
                    current_images = []
                    current_tables = []
                else:
                    current_content += para_text
            
            # Add final page
            if current_content.strip():
                pages.append({
                    'page_number': current_page,
                    'content': current_content.strip(),
                    'images': current_images,
                    'tables': current_tables,
                    'has_text': len(current_content.strip()) > 0,
                    'character_count': len(current_content)
                })
            
            # Distribute images and tables across pages
            images_per_page = len(images) // len(pages) if pages else 0
            tables_per_page = len(tables) // len(pages) if pages else 0
            
            for i, page in enumerate(pages):
                # Distribute images
                start_img = i * images_per_page
                end_img = start_img + images_per_page if i < len(pages) - 1 else len(images)
                page['images'] = images[start_img:end_img]
                
                # Distribute tables
                start_tbl = i * tables_per_page
                end_tbl = start_tbl + tables_per_page if i < len(pages) - 1 else len(tables)
                page['tables'] = tables[start_tbl:end_tbl]
        
        return pages
    
    def _is_heading_style(self, style_name: str) -> bool:
        """Check if paragraph style is a heading
        
        Args:
            style_name: Style name
            
        Returns:
            True if heading style
        """
        heading_styles = [
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
            'Title', 'Subtitle'
        ]
        return style_name in heading_styles
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions
        
        Returns:
            List of supported extensions
        """
        return ['.docx', '.doc']